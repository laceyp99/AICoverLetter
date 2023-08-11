import docx
import os
import platform
import openai


def generate_cover_letter(job_description, resume):
    # Set your OpenAI API key here
    openai.api_key = 'ADD YOUR API KEY HERE'

    prompt = f"Write a cover letter for the following job description:\n{job_description}\n\nHere is my resume:\n{resume}\n\n"

    print("ChatGPT Prompt:\n", prompt)

    response = openai.Completion.create(
        engine="text-davinci-003",  # Use DaVinci model
        prompt=prompt,
        max_tokens=1000 # Alter this to limit number of tokens used
        # stop=["Sincerely,"]
    )

    print("ChatGPT Response:\n", response)
    cover_letter = response.choices[0].text.strip()

    return cover_letter

def create_word_document(cover_letter):
    custom_header = "ADD A CUSTOM HEADER HERE"
    final_cover_letter = custom_header + cover_letter
    doc = docx.Document()
    doc.add_paragraph(final_cover_letter)

    doc.save("cover_letter.docx")

def convert_to_pdf(pdf_filename):
    if platform.system() == 'Windows':
        import win32com.client
        wdFormatPDF = 17
        word = win32com.client.Dispatch('Word.Application')
        doc_path = os.path.abspath("cover_letter.docx")

        # Save the PDF in the "exports" subfolder
        pdf_path = os.path.abspath(os.path.join("Cover Letters", pdf_filename + ".pdf"))

        doc = word.Documents.Open(doc_path)
        doc.SaveAs(pdf_path, FileFormat=wdFormatPDF)
        doc.Close()
        word.Quit()
    else:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph

        doc_path = "cover_letter.docx"
        pdf_path = "cover_letter.pdf"

        # Load the cover letter from Word document
        doc = docx.Document(doc_path)
        cover_letter = ""
        for paragraph in doc.paragraphs:
            cover_letter += paragraph.text + "\n"

        # Create PDF using ReportLab
        doc = SimpleDocTemplate(pdf_path, pagesize=letter)
        styles = getSampleStyleSheet()
        content = [Paragraph(cover_letter, styles["Normal"])]
        doc.build(content)

if __name__ == "__main__":
    # Read the job description from the file
    job_description_file = "job_description.txt"
    with open(job_description_file, "r", encoding="utf-8") as f:
        job_description = f.read()

    # Read the resume from the file
    resume_file = "resume.txt"
    with open(resume_file, "r", encoding="utf-8") as f:
        resume = f.read()

    
    # Generate the cover letter using the provided job description and resume
    cover_letter = generate_cover_letter(job_description, resume)
    print("\nGenerated Cover Letter:\n", cover_letter)
    

    create_word_document(cover_letter)
    print("Cover letter saved as 'cover_letter.docx'.")

    confirm_export_to_pdf = input("Do you want to export the cover letter to PDF? (yes/no): ")
    if confirm_export_to_pdf.lower() == "yes":
        # Prompt user to enter the desired PDF filename
        pdf_filename = input("Enter the PDF filename (without extension): ")
        convert_to_pdf(pdf_filename)
        print(f"Cover letter exported to PDF as '{pdf_filename}.pdf'.")
